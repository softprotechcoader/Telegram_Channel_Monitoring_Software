import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QFileDialog, QLabel, QVBoxLayout, QHBoxLayout, QWidget, QInputDialog, QScrollArea, QTextEdit, QProgressBar
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QIcon
from excel_handler import generate_empty_excel, load_excel
import asyncio
from telegram_handler import main as telegram_main
from qasync import QEventLoop, QApplication as QAsyncApplication
import pandas as pd

class TelegramAutomationApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Telegram Automation")
        self.setGeometry(100, 100, 800, 600)  # Adjusted for responsiveness

        # Set Telegram Automation icon
        self.setWindowIcon(QIcon("telegram_icon.png"))

        # Layout and widgets
        main_layout = QVBoxLayout()

        # Row 1: Generate Empty Excel and Load Excel File
        row1_layout = QHBoxLayout()
        self.generate_button = QPushButton("Generate Empty Excel")
        self.generate_button.setFixedSize(150, 30)
        self.generate_button.setStyleSheet("margin-left: 10px;")
        row1_layout.addWidget(self.generate_button, alignment=Qt.AlignLeft)

        self.load_button = QPushButton("Load Excel File")
        self.load_button.setFixedSize(150, 30)
        self.load_button.setStyleSheet("margin-right: 10px;")
        row1_layout.addWidget(self.load_button, alignment=Qt.AlignRight)

        main_layout.addLayout(row1_layout)

        # Row 2: Start Automation (centered)
        self.start_button = QPushButton("Start Automation")
        self.start_button.setFixedSize(200, 40)  # Medium size
        self.start_button.setStyleSheet("font-size: 14px; font-weight: bold;")
        main_layout.addWidget(self.start_button, alignment=Qt.AlignCenter)

        # Row 3: No of latest content and Read button
        row3_layout = QHBoxLayout()
        row3_layout.setSpacing(0)  # Remove spacing between widgets

        self.latest_content_label = QLabel("No of latest content:")
        self.latest_content_label.setStyleSheet("margin-left: 10px;")
        row3_layout.addWidget(self.latest_content_label, alignment=Qt.AlignLeft)

        self.latest_content_input = QTextEdit()
        self.latest_content_input.setFixedHeight(30)
        self.latest_content_input.setFixedWidth(250)
        row3_layout.addWidget(self.latest_content_input, alignment=Qt.AlignLeft)

        self.read_button = QPushButton("Read")
        self.read_button.setFixedSize(100, 30)
        row3_layout.addWidget(self.read_button, alignment=Qt.AlignLeft)

        main_layout.addLayout(row3_layout)

        # Row 4: Abort and Export to Word
        row4_layout = QHBoxLayout()
        self.abort_button = QPushButton("Abort")
        self.abort_button.setFixedSize(100, 30)
        self.abort_button.setStyleSheet("margin-left: 10px;")
        row4_layout.addWidget(self.abort_button, alignment=Qt.AlignLeft)

        self.export_button = QPushButton("Export to Word")
        self.export_button.setFixedSize(150, 30)
        self.export_button.setStyleSheet("margin-right: 10px;")
        row4_layout.addWidget(self.export_button, alignment=Qt.AlignRight)

        main_layout.addLayout(row4_layout)

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        main_layout.addWidget(self.progress_bar)

        # Status label
        self.label = QLabel("Status: Ready")
        self.label.setStyleSheet("font-size: 12px; color: blue;")
        main_layout.addWidget(self.label)

        # Scrollable area for displaying content
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.scroll_area.setWidget(self.text_display)
        main_layout.addWidget(self.scroll_area)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

        self.accounts = None
        self.abort_flag = False
        self.channel_content = []

        # Connect buttons to their respective functions
        self.generate_button.clicked.connect(self.generate_excel)
        self.load_button.clicked.connect(self.load_excel_file)
        self.start_button.clicked.connect(self.start_automation)
        self.abort_button.clicked.connect(self.abort_process)
        self.export_button.clicked.connect(self.export_to_word)
        self.read_button.clicked.connect(self.read_latest_content)

    def generate_excel(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Excel File", "", "Excel Files (*.xlsx)", options=options)
        if file_path:
            generate_empty_excel(file_path)
            self.label.setText(f"Empty Excel file generated at: {file_path}")

    def load_excel_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Excel File", "", "Excel Files (*.xlsx)", options=options)
        if file_path:
            try:
                self.accounts = load_excel(file_path).to_dict(orient='records')
                self.label.setText(f"Loaded Excel file: {file_path}")
            except ValueError as e:
                self.label.setText(str(e))

    def start_automation(self):
        if self.accounts:
            self.label.setText("Starting Telegram automation...")
            self.abort_flag = False
            loop = asyncio.get_event_loop()
            loop.create_task(self.run_telegram_automation())
        else:
            self.label.setText("Please load a valid Excel file first.")

    def abort_process(self):
        self.abort_flag = True
        self.label.setText("Process aborted. Resetting tool...")

        # Reset UI components but keep the loaded accounts
        self.progress_bar.setValue(0)
        self.text_display.clear()
        self.channel_content = []
        self.label.setText("Tool reset. Ready for new actions.")

    def export_to_word(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Word File", "", "Word Files (*.docx)", options=options)
        if file_path:
            from docx import Document

            # Create a new Word document
            doc = Document()
            doc.add_heading("Telegram Channel Content", level=1)

            # Organize content by date and channel
            data_by_date = {}
            for entry in self.channel_content:
                date = entry["Date"].date() if isinstance(entry["Date"], pd.Timestamp) else entry["Date"]
                channel_name = entry["Channel Name"]
                content = entry["Content"]

                if date not in data_by_date:
                    data_by_date[date] = {}
                data_by_date[date][channel_name] = content

            # Write data to the Word document
            for date, channels in sorted(data_by_date.items()):
                doc.add_heading(str(date), level=2)
                for channel, content in channels.items():
                    doc.add_paragraph(f"Channel: {channel}")
                    doc.add_paragraph(f"Content: {content}")
                    doc.add_paragraph("---")

            # Save the Word document
            doc.save(file_path)
            self.label.setText(f"Content exported to: {file_path}")

    def read_latest_content(self):
        try:
            n = int(self.latest_content_input.toPlainText().strip())
            if n <= 0:
                self.label.setText("Please enter a positive number.")
                return
            self.abort_flag = False  # Reset abort flag
            self.label.setText(f"Reading last {n} contents from all channels...")
            loop = asyncio.get_event_loop()
            loop.create_task(self.run_telegram_automation(last_n=n))
        except ValueError:
            self.label.setText("Invalid input. Please enter a valid number.")

    async def run_telegram_automation(self, last_n=None):
        async def get_login_code(phone_number):
            code, ok = QInputDialog.getText(self, "Login Code", f"Enter the login code for {phone_number}:")
            if ok and code:
                return code
            else:
                return None

        async def monitor_channels_with_ui(client):
            total_channels = 0
            processed_channels = 0

            async for dialog in client.iter_dialogs():
                if dialog.is_channel:
                    total_channels += 1

            async for dialog in client.iter_dialogs():
                if dialog.is_channel:
                    self.text_display.append(f"Monitoring channel: {dialog.name}")
                    async for message in client.iter_messages(dialog.id, limit=last_n):
                        if self.abort_flag:
                            return
                        content = f"[{dialog.name}] {message.date}: {message.text}"
                        self.text_display.append(content)
                        self.channel_content.append({"Date": message.date, "Channel Name": dialog.name, "Content": message.text})
                    processed_channels += 1
                    self.progress_bar.setValue(int((processed_channels / total_channels) * 100))

        await telegram_main(self.accounts, get_login_code, monitor_channels_with_ui)
        self.label.setText("Work completed successfully!")

if __name__ == "__main__":
    app = QAsyncApplication(sys.argv)
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)

    window = TelegramAutomationApp()
    window.show()

    with loop:
        loop.run_forever()